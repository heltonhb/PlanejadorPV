"""
Módulo para processamento e extração de texto de documentos.
"""

import logging
import os
import re
import tempfile
import traceback
from pathlib import Path
from typing import Optional

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import pdfplumber
import fitz
from chromadb import PersistentClient
from chromadb.utils.embedding_functions import DefaultEmbeddingFunction
from langchain_text_splitters import RecursiveCharacterTextSplitter

CHROMA_PATH = Path("vector_store")
COLLECTION_NAME = "documentos_ensina_mais"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

TESSDATA_PATH = Path("tessdata")
TESSDATA_URL = "https://github.com/tesseract-ocr/tessdata/raw/main/por.traineddata"


def _baixar_modelo_ocr(caminho: Path) -> None:
    """Baixa o modelo OCR por.traineddata usando requests (não usa signal.signal)."""
    import requests as req
    logger.info("Baixando modelo OCR português...")
    r = req.get(TESSDATA_URL, timeout=30)
    r.raise_for_status()
    caminho.write_bytes(r.content)
    logger.info("Modelo OCR baixado: %s", caminho)

EXTENSOES_SUPORTADAS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "texto",
    ".md": "texto",
    ".png": "imagem",
    ".jpg": "imagem",
    ".jpeg": "imagem",
    ".gif": "imagem",
    ".bmp": "imagem",
    ".tiff": "imagem",
}


def sanitizar_id(nome: str) -> str:
    """Sanitiza um nome para uso como ID, removendo caracteres especiais."""
    return re.sub(r'[^a-zA-Z0-9_\-.:]', '_', nome)[:120]


def _get_collection():
    """Obtém ou cria a collection do ChromaDB."""
    CHROMA_PATH.mkdir(exist_ok=True)
    client = PersistentClient(str(CHROMA_PATH))
    return client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=DefaultEmbeddingFunction(),
    )


def _extrair_pdfplumber(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de PDF usando pdfplumber."""
    try:
        with pdfplumber.open(tmp_path) as pdf:
            paginas = len(pdf.pages)
            partes = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    partes.append(t)
            return "\n".join(partes), paginas
    except Exception:
        logger.warning("pdfplumber falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_pymupdf(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de PDF usando PyMuPDF."""
    try:
        doc = fitz.open(tmp_path)
        paginas = len(doc)
        texto = "\n".join(page.get_text() or "" for page in doc)
        doc.close()
        return texto, paginas
    except Exception:
        logger.warning("pymupdf falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_ocr(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de PDF usando OCR (Tesseract)."""
    try:
        from PIL import Image
        from tesserocr import PyTessBaseAPI
    except ImportError as e:
        logger.warning("tesserocr não disponível: %s", e)
        return None, 0

    try:
        TESSDATA_PATH.mkdir(exist_ok=True)
        traineddata = TESSDATA_PATH / "por.traineddata"
        if not traineddata.exists():
            _baixar_modelo_ocr(traineddata)
        
        doc = fitz.open(tmp_path)
        paginas = len(doc)
        api = PyTessBaseAPI(lang="por", path=str(TESSDATA_PATH))
        textos = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            api.SetImage(img)
            t = api.GetUTF8Text()
            if t:
                textos.append(t.strip())
        api.End()
        doc.close()
        return "\n\n".join(textos), paginas
    except Exception:
        logger.warning("ocr falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_docx(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de documentos DOCX."""
    try:
        from docx import Document
        
        doc = Document(tmp_path)
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        texto = "\n\n".join(paragrafos)
        
        tabelas_texto = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    tabelas_texto.append(row_text)
        
        if tabelas_texto:
            texto += "\n\n[Tabelas]\n" + "\n".join(tabelas_texto)
        
        return texto, len(paragrafos)
    except ImportError:
        logger.warning("python-docx não disponível")
        return None, 0
    except Exception:
        logger.warning("docx falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_texto_puro(tmp_path: str) -> tuple[str, int]:
    """Extrai texto de arquivos de texto puro."""
    try:
        with open(tmp_path, "r", encoding="utf-8") as f:
            texto = f.read()
        linhas = texto.split("\n")
        return texto, len(linhas)
    except Exception:
        logger.warning("texto puro falhou:\n%s", traceback.format_exc())
        return "", 0


def _extrair_imagem_ocr(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de imagens usando OCR."""
    try:
        from PIL import Image
        from tesserocr import PyTessBaseAPI
    except ImportError as e:
        logger.warning("tesserocr/PIL não disponível: %s", e)
        return None, 0

    try:
        TESSDATA_PATH.mkdir(exist_ok=True)
        traineddata = TESSDATA_PATH / "por.traineddata"
        if not traineddata.exists():
            _baixar_modelo_ocr(traineddata)
        img = Image.open(tmp_path)
        api = PyTessBaseAPI(lang="por", path=str(TESSDATA_PATH))
        api.SetImage(img)
        texto = api.GetUTF8Text()
        api.End()
        img.close()
        return texto.strip(), 1
    except Exception:
        logger.warning("imagem OCR falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_texto_worker(pdf_bytes: bytes) -> dict:
    """Worker function para extração de texto de PDF."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    
    erros = []
    try:
        texto, paginas = _extrair_pdfplumber(tmp_path)
        metodo = "pdfplumber"
        if not texto or not texto.strip():
            logger.info("pdfplumber não extraiu texto, tentando pymupdf")
            texto, paginas = _extrair_pymupdf(tmp_path)
            metodo = "pymupdf"
        if not texto or not texto.strip():
            logger.info("pymupdf não extraiu texto, tentando OCR")
            texto, paginas = _extrair_ocr(tmp_path)
            metodo = "ocr"
        
        if not texto or not texto.strip():
            logger.warning("nenhum método extraiu texto do PDF (%d páginas)", paginas)
        
        return {"texto": texto or "", "metodo": metodo, "paginas": paginas}
    except Exception as e:
        logger.error("extração falhou: %s\n%s", e, traceback.format_exc())
        return {"texto": "", "metodo": f"erro: {e}", "paginas": 0}
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def extrair_texto(pdf_bytes: bytes) -> dict:
    """Extrai texto de bytes de PDF."""
    return _extrair_texto_worker(pdf_bytes)


def extrair_arquivo(bytes_arquivo: bytes, nome_arquivo: str) -> dict:
    """Extrai texto de um arquivo baseado em sua extensão."""
    extensao = Path(nome_arquivo).suffix.lower()
    tipo = EXTENSOES_SUPORTADAS.get(extensao, "desconhecido")
    
    if tipo == "desconhecido":
        return {
            "texto": "",
            "metodo": "extensão não suportada",
            "paginas": 0,
            "erro": f"Extensão {extensao} não suportada",
        }
    
    with tempfile.NamedTemporaryFile(suffix=extensao, delete=False) as tmp:
        tmp.write(bytes_arquivo)
        tmp_path = tmp.name
    
    try:
        if tipo == "pdf":
            return extrair_texto(bytes_arquivo)
        
        elif tipo == "docx":
            texto, paginas = _extrair_docx(tmp_path)
            return {
                "texto": texto or "",
                "metodo": "docx",
                "paginas": paginas,
            }
        
        elif tipo == "texto":
            texto, paginas = _extrair_texto_puro(tmp_path)
            return {
                "texto": texto,
                "metodo": "texto",
                "paginas": paginas,
            }
        
        elif tipo == "imagem":
            texto, paginas = _extrair_imagem_ocr(tmp_path)
            return {
                "texto": texto or "",
                "metodo": "ocr",
                "paginas": paginas,
            }
        
        return {"texto": "", "metodo": "erro", "paginas": 0}
    
    finally:
        os.unlink(tmp_path)


def chunk_texto(texto: str) -> list[dict]:
    """Divide texto em chunks para indexação."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    chunks = splitter.create_documents([texto])
    return [
        {
            "texto": chunk.page_content,
            "metadata": chunk.metadata or {},
        }
        for chunk in chunks
    ]


def salvar_chunks(
    chunks: list[dict],
    documento_id: str = None,
    extra_metadata: dict = None,
) -> int:
    """Salva chunks no ChromaDB e opcionalmente no Firebase."""
    collection = _get_collection()
    ids = []
    textos = []
    metadatas = []
    doc_id = sanitizar_id(documento_id or "doc")

    for i, chunk in enumerate(chunks):
        chunk_id = f"{doc_id}_{i}"
        ids.append(chunk_id)
        textos.append(chunk["texto"])
        md = {**chunk["metadata"], "chunk_id": chunk_id}
        if extra_metadata:
            md.update(extra_metadata)
        md["documento_id"] = doc_id
        metadatas.append(md)

    collection.upsert(documents=textos, metadatas=metadatas, ids=ids)

    try:
        from utils.firebase_store import salvar_chunks_firestore
        salvar_chunks_firestore(ids, textos, metadatas)
    except Exception:
        logger.debug("Firebase não disponível para salvar (ignorado)")

    return len(chunks)


def salvar_resumo_documento(documento_id: str, resumo: str):
    """Atualiza o campo 'resumo' em todos os chunks de um documento."""
    collection = _get_collection()
    try:
        resultados = collection.get(where={"documento_id": sanitizar_id(documento_id)})
        if not resultados or not resultados["ids"]:
            return
        novos_metadatas = []
        for md in resultados["metadatas"]:
            md["resumo"] = resumo
            novos_metadatas.append(md)
        collection.update(
            ids=resultados["ids"],
            metadatas=novos_metadatas,
        )
    except Exception as e:
        logger.debug(f"Erro ao salvar resumo no ChromaDB: {e}")


def processar_documento(pdf_bytes: bytes, nome_arquivo: str = None) -> dict:
    """Processa um documento PDF: extrai texto, cria chunks e salva."""
    resultado_extracao = extrair_texto(pdf_bytes)
    texto = resultado_extracao["texto"]
    paginas = resultado_extracao["paginas"]
    metodo = resultado_extracao["metodo"]

    if not texto.strip():
        msg = f"Nenhum texto extraído do PDF ({paginas} páginas, método: {metodo})."
        if metodo.startswith("erro:"):
            msg += f" Detalhes: {metodo[5:]}"
        elif metodo == "ocr":
            msg += " O PDF pode ser um documento escaneado (imagem) e o OCR não está disponível neste ambiente."
        return {
            "status": "erro",
            "mensagem": msg,
            "paginas": paginas,
            "metodo": metodo,
        }
    
    chunks = chunk_texto(texto)
    total = salvar_chunks(
        chunks,
        documento_id=nome_arquivo,
        extra_metadata={"fonte": "pdf", "arquivo": nome_arquivo},
    )
    return {
        "status": "ok",
        "total_chunks": total,
        "total_caracteres": len(texto),
        "paginas": paginas,
        "metodo": metodo,
        "texto_completo": texto,
    }


def processar_arquivo(bytes_arquivo: bytes, nome_arquivo: str) -> dict:
    """Processa um arquivo qualquer: extrai texto, cria chunks e salva."""
    resultado = extrair_arquivo(bytes_arquivo, nome_arquivo)
    texto = resultado["texto"]
    paginas = resultado["paginas"]
    metodo = resultado["metodo"]
    
    if "erro" in resultado:
        return {
            "status": "erro",
            "mensagem": resultado["erro"],
            "paginas": 0,
            "metodo": metodo,
        }

    if not texto.strip():
        return {
            "status": "erro",
            "mensagem": f"Nenhum texto extraído ({nome_arquivo}, método: {metodo}).",
            "paginas": paginas,
            "metodo": metodo,
        }
    
    chunks = chunk_texto(texto)
    extensao = Path(nome_arquivo).suffix.lower()
    total = salvar_chunks(
        chunks,
        documento_id=nome_arquivo,
        extra_metadata={"fonte": extensao, "arquivo": nome_arquivo},
    )
    return {
        "status": "ok",
        "total_chunks": total,
        "total_caracteres": len(texto),
        "paginas": paginas,
        "metodo": metodo,
        "texto_completo": texto,
    }
