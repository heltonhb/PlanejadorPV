import csv
import io
import re
from datetime import datetime


# ---------------------------------------------------------------------------
# CSV exports
# ---------------------------------------------------------------------------

def exportar_relatorio_csv(relatorio: dict) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Fonte", "Título", "Chunks", "Caracteres", "Resumo"])
    for item in relatorio.get("fontes_detalhadas", []):
        writer.writerow([
            item.get("fonte", ""),
            item.get("titulo", ""),
            item.get("chunks", 0),
            item.get("caracteres", 0),
            item.get("resumo", ""),
        ])
    return output.getvalue()


def exportar_calendario_csv(conteudo_md: str, mes: str, ano: int) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Data", "Tipo", "Conteudo"])
    current_date = ""
    current_type = ""
    for line in conteudo_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        date_match = re.match(r"^\*\*(\d{1,2}/\d{1,2})\*\*", line)
        if date_match:
            current_date = date_match.group(1)
            current_type = ""
            rest = line[date_match.end():].strip().lstrip("-:").strip()
            if rest:
                writer.writerow([current_date, "", rest])
        elif line.startswith("### ") or line.startswith("## "):
            current_type = line.lstrip("#").strip()
        elif current_date:
            writer.writerow([current_date, current_type, line.lstrip("-").strip()])
    return output.getvalue()


def exportar_campanha_csv(dados: dict, conteudo_md: str) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Campo", "Valor"])
    if dados.get("nome"):
        writer.writerow(["Nome", dados["nome"]])
    if dados.get("objetivo"):
        writer.writerow(["Objetivo", dados["objetivo"]])
    if dados.get("publico"):
        writer.writerow(["Publico-alvo", dados["publico"]])
    if dados.get("servico"):
        writer.writerow(["Servico", dados["servico"]])
    if dados.get("orcamento"):
        writer.writerow(["Orcamento", f"R$ {dados['orcamento']:,.2f}"])
    if dados.get("canais"):
        writer.writerow(["Canais", ", ".join(dados["canais"])])
    if dados.get("datas"):
        writer.writerow(["Periodo", dados["datas"]])
    writer.writerow([])
    writer.writerow(["Secao", "Conteudo"])
    current_section = ""
    for line in conteudo_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            current_section = line[4:].strip()
        elif line.startswith("## "):
            current_section = line[3:].strip()
        elif line.startswith("# "):
            current_section = line[2:].strip()
        elif current_section:
            writer.writerow([current_section, line.lstrip("-").strip()])
            current_section = ""
    return output.getvalue()


# ---------------------------------------------------------------------------
# XLSX exports (openpyxl)
# ---------------------------------------------------------------------------

def exportar_relatorio_xlsx(relatorio: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    # -- Summary sheet --
    ws_summary = wb.active
    ws_summary.title = "Resumo"
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="00A859", end_color="00A859", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    summary_headers = ["Metrica", "Valor"]
    for col, header in enumerate(summary_headers, 1):
        cell = ws_summary.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    summary_data = [
        ("Total de Trechos", relatorio.get("total_chunks", 0)),
        ("Total de Caracteres", relatorio.get("total_caracteres", 0)),
        ("Tipos de Fonte", len(relatorio.get("por_fonte", {}))),
    ]
    for row, (metric, value) in enumerate(summary_data, 2):
        ws_summary.cell(row=row, column=1, value=metric).border = thin_border
        ws_summary.cell(row=row, column=2, value=value).border = thin_border

    ws_summary.column_dimensions["A"].width = 25
    ws_summary.column_dimensions["B"].width = 20

    # -- Detailed sources sheet --
    ws_detail = wb.create_sheet("Fontes Detalhadas")
    detail_headers = ["Fonte", "Titulo", "Chunks", "Caracteres", "Resumo"]
    for col, header in enumerate(detail_headers, 1):
        cell = ws_detail.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    for row, item in enumerate(relatorio.get("fontes_detalhadas", []), 2):
        ws_detail.cell(row=row, column=1, value=item.get("fonte", "")).border = thin_border
        ws_detail.cell(row=row, column=2, value=item.get("titulo", "")).border = thin_border
        ws_detail.cell(row=row, column=3, value=item.get("chunks", 0)).border = thin_border
        ws_detail.cell(row=row, column=4, value=item.get("caracteres", 0)).border = thin_border
        ws_detail.cell(row=row, column=5, value=item.get("resumo", "")).border = thin_border

    ws_detail.column_dimensions["A"].width = 15
    ws_detail.column_dimensions["B"].width = 45
    ws_detail.column_dimensions["C"].width = 12
    ws_detail.column_dimensions["D"].width = 15
    ws_detail.column_dimensions["E"].width = 60

    # -- Source distribution sheet --
    ws_dist = wb.create_sheet("Distribuicao")
    dist_headers = ["Tipo de Fonte", "Chunks", "Caracteres"]
    for col, header in enumerate(dist_headers, 1):
        cell = ws_dist.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    for row, (fonte, dados) in enumerate(sorted(relatorio.get("por_fonte", {}).items()), 2):
        ws_dist.cell(row=row, column=1, value=fonte).border = thin_border
        ws_dist.cell(row=row, column=2, value=dados.get("chunks", 0)).border = thin_border
        ws_dist.cell(row=row, column=3, value=dados.get("caracteres", 0)).border = thin_border

    ws_dist.column_dimensions["A"].width = 20
    ws_dist.column_dimensions["B"].width = 12
    ws_dist.column_dimensions["C"].width = 15

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def exportar_calendario_xlsx(conteudo_md: str, mes: str, ano: int) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = f"Calendario {mes} {ano}"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="00A859", end_color="00A859", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    headers = ["Data", "Secao", "Conteudo"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    current_date = ""
    current_section = ""
    row_num = 2

    for line in conteudo_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        date_match = re.match(r"^\*\*(\d{1,2}/\d{1,2})\*\*", line)
        if date_match:
            current_date = date_match.group(1)
            current_section = ""
            rest = line[date_match.end():].strip().lstrip("-:").strip()
            if rest:
                ws.cell(row=row_num, column=1, value=current_date).border = thin_border
                ws.cell(row=row_num, column=2, value="").border = thin_border
                ws.cell(row=row_num, column=3, value=rest).border = thin_border
                row_num += 1
        elif line.startswith("### "):
            current_section = line[4:].strip()
        elif line.startswith("## "):
            current_section = line[3:].strip()
        elif current_date:
            ws.cell(row=row_num, column=1, value=current_date).border = thin_border
            ws.cell(row=row_num, column=2, value=current_section).border = thin_border
            ws.cell(row=row_num, column=3, value=line.lstrip("-").strip()).border = thin_border
            row_num += 1

    ws.column_dimensions["A"].width = 15
    ws.column_dimensions["B"].width = 25
    ws.column_dimensions["C"].width = 70

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def exportar_campanha_xlsx(dados: dict, conteudo_md: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="00A859", end_color="00A859", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    # -- Metadata sheet --
    ws_meta = wb.active
    ws_meta.title = "Dados da Campanha"
    meta_headers = ["Campo", "Valor"]
    for col, header in enumerate(meta_headers, 1):
        cell = ws_meta.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    meta_rows = [
        ("Nome", dados.get("nome", "")),
        ("Objetivo", dados.get("objetivo", "")),
        ("Publico-alvo", dados.get("publico", "")),
        ("Servico", dados.get("servico", "")),
        ("Orcamento", f"R$ {dados.get('orcamento', 0):,.2f}"),
        ("Canais", ", ".join(dados.get("canais", []))),
        ("Periodo", dados.get("datas", "")),
    ]
    for row, (campo, valor) in enumerate(meta_rows, 2):
        ws_meta.cell(row=row, column=1, value=campo).border = thin_border
        ws_meta.cell(row=row, column=2, value=valor).border = thin_border

    ws_meta.column_dimensions["A"].width = 20
    ws_meta.column_dimensions["B"].width = 60

    # -- Content sheet --
    ws_content = wb.create_sheet("Conteudo")
    content_headers = ["Secao", "Conteudo"]
    for col, header in enumerate(content_headers, 1):
        cell = ws_content.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    current_section = ""
    row_num = 2
    for line in conteudo_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            current_section = line[4:].strip()
        elif line.startswith("## "):
            current_section = line[3:].strip()
        elif line.startswith("# "):
            current_section = line[2:].strip()
        elif current_section:
            ws_content.cell(row=row_num, column=1, value=current_section).border = thin_border
            ws_content.cell(row=row_num, column=2, value=line.lstrip("-").strip()).border = thin_border
            row_num += 1

    ws_content.column_dimensions["A"].width = 30
    ws_content.column_dimensions["B"].width = 80

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF exports (fpdf2)
# ---------------------------------------------------------------------------

def _criar_pdf_base(titulo: str) -> "FPDF":
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(0, 168, 89)
            self.cell(0, 8, "PlanejadorPV — Marketing Inteligente", new_x="LMARGIN", new_y="NEXT", align="R")
            self.set_draw_color(0, 168, 89)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Pagina {self.page_no()}/{{nb}}", align="C")

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 12, titulo, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(2)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    return pdf


def _adicionar_md_ao_pdf(pdf: "FPDF", conteudo_md: str):
    for line in conteudo_md.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        if line.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(4)
            pdf.multi_cell(0, 8, line[2:].strip())
            pdf.ln(2)
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(0, 100, 50)
            pdf.ln(3)
            pdf.multi_cell(0, 7, line[3:].strip())
            pdf.ln(1)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 80, 40)
            pdf.ln(2)
            pdf.multi_cell(0, 6, line[4:].strip())
            pdf.ln(1)
        else:
            clean = line.lstrip("-").strip()
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", clean)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(40, 40, 40)
            prefix = "  •  " if line.startswith("-") else ""
            pdf.multi_cell(0, 5, prefix + clean)


def exportar_relatorio_pdf(relatorio: dict) -> bytes:
    pdf = _criar_pdf_base("Relatorio de Conteudo")

    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Resumo Geral", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, f"Total de trechos: {relatorio.get('total_chunks', 0)}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Total de caracteres: {relatorio.get('total_caracteres', 0):,}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Tipos de fonte: {len(relatorio.get('por_fonte', {}))}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Distribuicao por Tipo de Fonte", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for fonte, dados in sorted(relatorio.get("por_fonte", {}).items()):
        pdf.cell(0, 6, f"  {fonte.upper()}: {dados.get('chunks', 0)} trechos, {dados.get('caracteres', 0):,} caracteres", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Fontes Detalhadas", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    for item in relatorio.get("fontes_detalhadas", []):
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(0, 80, 40)
        pdf.cell(0, 6, f"{item.get('titulo', 'Sem titulo')}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 5, f"Fonte: {item.get('fonte', '')} | Trechos: {item.get('chunks', 0)} | Caracteres: {item.get('caracteres', 0):,}", new_x="LMARGIN", new_y="NEXT")
        if item.get("resumo"):
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(60, 60, 60)
            pdf.multi_cell(0, 5, f"Resumo: {item['resumo']}")
        pdf.ln(3)

    return bytes(pdf.output())


def exportar_calendario_pdf(conteudo_md: str, mes: str, ano: int) -> bytes:
    pdf = _criar_pdf_base(f"Calendario Editorial — {mes} {ano}")
    _adicionar_md_ao_pdf(pdf, conteudo_md)
    return bytes(pdf.output())


def exportar_campanha_pdf(dados: dict, conteudo_md: str) -> bytes:
    pdf = _criar_pdf_base("Campanha de Marketing")

    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Dados da Campanha", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    if dados.get("nome"):
        pdf.cell(0, 6, f"Nome: {dados['nome']}", new_x="LMARGIN", new_y="NEXT")
    if dados.get("objetivo"):
        pdf.cell(0, 6, f"Objetivo: {dados['objetivo']}", new_x="LMARGIN", new_y="NEXT")
    if dados.get("publico"):
        pdf.cell(0, 6, f"Publico-alvo: {dados['publico']}", new_x="LMARGIN", new_y="NEXT")
    if dados.get("servico"):
        pdf.cell(0, 6, f"Servico: {dados['servico']}", new_x="LMARGIN", new_y="NEXT")
    if dados.get("orcamento"):
        pdf.cell(0, 6, f"Orcamento: R$ {dados['orcamento']:,.2f}", new_x="LMARGIN", new_y="NEXT")
    if dados.get("canais"):
        pdf.cell(0, 6, f"Canais: {', '.join(dados['canais'])}", new_x="LMARGIN", new_y="NEXT")
    if dados.get("datas"):
        pdf.cell(0, 6, f"Periodo: {dados['datas']}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    _adicionar_md_ao_pdf(pdf, conteudo_md)
    return bytes(pdf.output())


# ---------------------------------------------------------------------------
# DOCX export (kept from original)
# ---------------------------------------------------------------------------

def exportar_markdown_docx(conteudo_md: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in conteudo_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, line[2:])
        elif line[0].isdigit() and ". " in line[:4]:
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, line.split(". ", 1)[1])
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_rich_text(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
