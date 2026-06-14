"""
Extração de texto de documentos: PDF, DOCX, imagens, texto puro.
Serve como camada de extração delegada pelo módulo documentos.py.
"""
import base64
import json
import logging
import os
import subprocess
import sys
import tempfile
import traceback
from pathlib import Path
from typing import Optional

import pdfplumber

try:
    import fitz
    _HAS_FITZ = True
except ImportError:
    _HAS_FITZ = False

from utils.ocr import encontrar_modelo_ocr

logger = logging.getLogger(__name__)

# ── Extensões suportadas ──────────────────────────────────────────────

EXTENSOES_SUPORTADAS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "texto",
    ".md": "texto",
    ".png": "imagem",
    ".jpg": "imagem",
    ".jpeg": "imagem",
    ".gif": "imagem",
    ".bmp": "imagem",
    ".tiff": "imagem",
}

# ── Extractores individuais ───────────────────────────────────────────


def _extrair_pdfplumber(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de PDF usando pdfplumber."""
    try:
        with pdfplumber.open(tmp_path) as pdf:
            paginas = len(pdf.pages)
            partes = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    partes.append(t)
            return "\n".join(partes), paginas
    except Exception:
        logger.warning("pdfplumber falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_pymupdf(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de PDF usando PyMuPDF."""
    if not _HAS_FITZ:
        logger.info("pymupdf indisponível (não instalado)")
        return None, 0
    try:
        doc = fitz.open(tmp_path)
        paginas = len(doc)
        texto = "\n".join(page.get_text() or "" for page in doc)
        doc.close()
        return texto, paginas
    except Exception:
        logger.warning("pymupdf falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_ocr(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de PDF usando OCR (Tesseract).

    Tenta tesserocr primeiro (nativo, rápido). Se falhar no import
    (ex: signal.signal em thread não-principal), usa pytesseract
    como fallback (subprocesso, sem signal).
    """
    from PIL import Image

    # Tentar tesserocr (C native, mas falha em thread não-principal)
    try:
        from tesserocr import PyTessBaseAPI

        modelo = encontrar_modelo_ocr()
        if modelo and _HAS_FITZ:
            doc = fitz.open(tmp_path)
            paginas = len(doc)
            api = PyTessBaseAPI(lang="por", path=str(modelo.parent))
            textos = []
            for page in doc:
                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                api.SetImage(img)
                t = api.GetUTF8Text()
                if t:
                    textos.append(t.strip())
            api.End()
            doc.close()
            if textos:
                return "\n\n".join(textos), paginas
    except Exception as e:
        logger.warning("tesserocr falhou: %s", e)

    # Fallback: pytesseract (subprocesso, funciona em qualquer thread)
    logger.info("usando pytesseract como fallback")
    if not _HAS_FITZ:
        logger.info("pymupdf indisponível para OCR")
        return None, 0
    try:
        import pytesseract

        doc = fitz.open(tmp_path)
        paginas = len(doc)
        textos = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            t = pytesseract.image_to_string(img, lang="por")
            if t and t.strip():
                textos.append(t.strip())
        doc.close()
        return "\n\n".join(textos), paginas
    except Exception as e:
        logger.warning("pytesseract falhou: %s", e)
        return None, 0


def _extrair_docx(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de documentos DOCX."""
    try:
        from docx import Document

        doc = Document(tmp_path)
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        texto = "\n\n".join(paragrafos)

        tabelas_texto = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    tabelas_texto.append(row_text)

        if tabelas_texto:
            texto += "\n\n[Tabelas]\n" + "\n".join(tabelas_texto)

        return texto, len(paragrafos)
    except ImportError:
        logger.warning("python-docx não disponível")
        return None, 0
    except Exception:
        logger.warning("docx falhou:\n%s", traceback.format_exc())
        return None, 0


def _extrair_texto_puro(tmp_path: str) -> tuple[str, int]:
    """Extrai texto de arquivos de texto puro."""
    try:
        with open(tmp_path, "r", encoding="utf-8") as f:
            texto = f.read()
        linhas = texto.split("\n")
        return texto, len(linhas)
    except Exception:
        logger.warning("texto puro falhou:\n%s", traceback.format_exc())
        return "", 0


def _extrair_imagem_ocr(tmp_path: str) -> tuple[Optional[str], int]:
    """Extrai texto de imagens usando OCR."""
    from PIL import Image

    # Tentar tesserocr
    try:
        from tesserocr import PyTessBaseAPI

        modelo = encontrar_modelo_ocr()
        if modelo:
            img = Image.open(tmp_path)
            api = PyTessBaseAPI(lang="por", path=str(modelo.parent))
            api.SetImage(img)
            texto = api.GetUTF8Text()
            api.End()
            img.close()
            if texto and texto.strip():
                return texto.strip(), 1
    except Exception as e:
        logger.warning("tesserocr imagem falhou: %s", e)

    # Fallback: pytesseract
    try:
        import pytesseract

        img = Image.open(tmp_path)
        texto = pytesseract.image_to_string(img, lang="por")
        img.close()
        return texto.strip(), 1
    except Exception as e:
        logger.warning("pytesseract imagem falhou: %s", e)
        return None, 0


# ── Worker e dispatcher ───────────────────────────────────────────────


def _extrair_texto_worker(pdf_bytes: bytes) -> dict:
    """Worker function para extração de texto de PDF."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name

    erros = []
    try:
        texto, paginas = _extrair_pdfplumber(tmp_path)
        metodo = "pdfplumber"
        if not texto or not texto.strip():
            logger.info("pdfplumber não extraiu texto, tentando pymupdf")
            texto, paginas = _extrair_pymupdf(tmp_path)
            metodo = "pymupdf"
        if not texto or not texto.strip():
            logger.info("pymupdf não extraiu texto, tentando OCR")
            texto, paginas = _extrair_ocr(tmp_path)
            metodo = "ocr"

        if not texto or not texto.strip():
            logger.warning("nenhum método extraiu texto do PDF (%d páginas)", paginas)

        return {"texto": texto or "", "metodo": metodo, "paginas": paginas}
    except Exception as e:
        logger.error("extração falhou: %s\n%s", e, traceback.format_exc())
        return {"texto": "", "metodo": f"erro: {e}", "paginas": 0}
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def extrair_texto(pdf_bytes: bytes) -> dict:
    """Extrai texto de bytes de PDF.

    Tenta extração inline primeiro. Se falhar com RuntimeError
    de signal (thread não-principal no Cloud), faz fallback
    para subprocesso dedicado.
    """
    resultado = _extrair_texto_worker(pdf_bytes)

    # Se método começar com "erro:", verificar se é signal
    if resultado["metodo"].startswith("erro:") and "signal" in resultado["metodo"]:
        logger.warning("extração inline falhou por signal, tentando subprocesso")
        try:
            resultado = _extrair_subprocesso(pdf_bytes)
        except Exception as e2:
            logger.error("subprocesso também falhou: %s", e2)
            return {
                "status": "erro",
                "mensagem": f"Extração falhou mesmo em subprocesso: {e2}",
                "paginas": 0,
                "metodo": "erro",
            }

    if resultado.get("status") is None:
        return resultado
    return resultado


def _extrair_subprocesso(pdf_bytes: bytes) -> dict:
    """Executa extração de texto em um subprocesso dedicado.

    Necessário no Streamlit Cloud onde a thread principal
    pode não estar disponível para signal.signal().
    """
    worker_path = Path(__file__).parent / "_extract_worker.py"
    payload = {"pdf_b64": base64.b64encode(pdf_bytes).decode("ascii")}

    proc = subprocess.run(
        [sys.executable, str(worker_path)],
        input=json.dumps(payload).encode(),
        capture_output=True,
        timeout=120,
    )

    if proc.returncode != 0:
        stderr = proc.stderr.decode(errors="replace")
        raise RuntimeError(f"worker retornou código {proc.returncode}: {stderr}")

    try:
        return json.loads(proc.stdout.decode())
    except json.JSONDecodeError as e:
        raise RuntimeError(f"resposta inválida do worker: {e}")


def extrair_arquivo(bytes_arquivo: bytes, nome_arquivo: str) -> dict:
    """Extrai texto de um arquivo baseado em sua extensão."""
    extensao = Path(nome_arquivo).suffix.lower()
    tipo = EXTENSOES_SUPORTADAS.get(extensao, "desconhecido")

    if tipo == "desconhecido":
        return {
            "texto": "",
            "metodo": "extensão não suportada",
            "paginas": 0,
            "erro": f"Extensão {extensao} não suportada",
        }

    with tempfile.NamedTemporaryFile(suffix=extensao, delete=False) as tmp:
        tmp.write(bytes_arquivo)
        tmp_path = tmp.name

    try:
        if tipo == "pdf":
            return extrair_texto(bytes_arquivo)

        elif tipo == "docx":
            texto, paginas = _extrair_docx(tmp_path)
            return {
                "texto": texto or "",
                "metodo": "docx",
                "paginas": paginas,
            }

        elif tipo == "texto":
            texto, paginas = _extrair_texto_puro(tmp_path)
            return {
                "texto": texto,
                "metodo": "texto",
                "paginas": paginas,
            }

        elif tipo == "imagem":
            texto, paginas = _extrair_imagem_ocr(tmp_path)
            return {
                "texto": texto or "",
                "metodo": "ocr",
                "paginas": paginas,
            }

        return {"texto": "", "metodo": "erro", "paginas": 0}

    finally:
        os.unlink(tmp_path)
